import pandas as pd
from docx import Document
import os
import shutil
from docx.shared import Pt

BASE_CV_PATH = "./../Required documents/{name} CV.docx"

def read_excel(file_path):
    # Read the Excel file
    df = pd.read_excel(file_path, dtype=str)
    return df

def replace_placeholder(paragraph, placeholder, replacement):
    if placeholder in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if placeholder in item.text:
                item.text = item.text.replace(placeholder, replacement)
                
def create_word_from_template(template_path, output_path, data):
    # Load the Word template
    doc = Document(template_path)
    
    for _, row in data.iterrows():
        # Create company folder
        os.mkdir(f"{output_path}/{row["Company"]}")

        # Make a copy of the template
        new_doc = Document(template_path)
        
        # Replace placeholders with actual data
        for paragraph in new_doc.paragraphs:
            replace_placeholder(paragraph, "{Position}", row["Position"])
            replace_placeholder(paragraph, "{Company}", row["Company"])
            replace_placeholder(paragraph, "{Paragraphs}", row["Paragraphs"])
            replace_placeholder(paragraph, "{Languages}", row["Languages"])
            replace_placeholder(paragraph, "{Date}", row["Date"])
            replace_placeholder(paragraph, "{Platform}", row["Platform"])

        # Save the new document
        new_doc.save(f"{output_path}/{row["Company"]}/{row["Company"]} Cover letter.docx")

        # Copy the CV
        shutil.copyfile(BASE_CV_PATH, f"{output_path}/{row["Company"]}/{name} CV.docx")

if __name__ == "__main__":
    # File paths
    excel_file_path = 'data.xlsx'
    word_template_path = 'Temp.docx'
    output_directory = './../Required documents/Company dedicated'
    
    # Read data from Excel
    data = read_excel(excel_file_path)
    
    # Create Word documents from template
    create_word_from_template(word_template_path, output_directory, data)
